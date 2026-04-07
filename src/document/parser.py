"""
文档解析模块
支持 PDF、DOCX、TXT 格式，解决中文乱码问题
"""
import re
from pathlib import Path
from typing import Union, BinaryIO
from dataclasses import dataclass
import logging

logger = logging.getLogger(__name__)


class ParseError(Exception):
    """文档解析异常"""
    pass


@dataclass
class ParsedDocument:
    """解析后的文档"""
    text: str
    filename: str
    file_type: str
    page_count: int = 1
    metadata: dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentParser:
    """文档解析器 - 支持 PDF/DOCX/TXT"""
    
    ALLOWED_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'text/plain': 'txt'
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        self._load_parsers()
    
    def _load_parsers(self):
        """延迟加载解析库"""
        self._pypdf_available = False
        self._pdfplumber_available = False
        self._docx_available = False
        
        try:
            from pypdf import PdfReader
            self._pypdf_available = True
        except ImportError:
            logger.warning("pypdf 未安装，PDF 解析将使用备用方案")
        
        try:
            import pdfplumber
            self._pdfplumber_available = True
        except ImportError:
            logger.warning("pdfplumber 未安装")
        
        try:
            from docx import Document
            self._docx_available = True
        except ImportError:
            logger.warning("python-docx 未安装，DOCX 解析不可用")
    
    def validate_file(self, file: Union[bytes, BinaryIO], content_type: str, filename: str) -> str:
        """验证文件格式和大小"""
        # 检查文件类型
        if content_type not in self.ALLOWED_TYPES:
            allowed = ', '.join(self.ALLOWED_TYPES.keys())
            raise ParseError(f"不支持的文件类型: {content_type}。仅支持: {allowed}")
        
        # 检查文件大小
        if isinstance(file, bytes):
            size = len(file)
        else:
            file.seek(0, 2)
            size = file.tell()
            file.seek(0)
        
        if size > self.MAX_FILE_SIZE:
            raise ParseError(f"文件大小超过限制: {size / 1024 / 1024:.1f}MB > 10MB")
        
        # 检查文件扩展名
        ext = Path(filename).suffix.lower()
        type_from_ext = self._get_type_from_extension(ext)
        
        if type_from_ext and type_from_ext != self.ALLOWED_TYPES[content_type]:
            logger.warning(f"文件扩展名与内容类型不匹配: {ext} vs {content_type}")
        
        return self.ALLOWED_TYPES[content_type]
    
    def _get_type_from_extension(self, ext: str) -> str:
        """从扩展名获取文件类型"""
        ext_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.txt': 'txt'
        }
        return ext_map.get(ext)
    
    def parse(self, file: Union[bytes, BinaryIO], content_type: str, filename: str) -> ParsedDocument:
        """
        解析文档
        
        Args:
            file: 文件内容或文件对象
            content_type: MIME 类型
            filename: 文件名
        
        Returns:
            ParsedDocument: 解析后的文档
        """
        file_type = self.validate_file(file, content_type, filename)
        
        # 统一转为 bytes
        if hasattr(file, 'read'):
            content = file.read()
        else:
            content = file
        
        try:
            if file_type == 'pdf':
                return self._parse_pdf(content, filename)
            elif file_type == 'docx':
                return self._parse_docx(content, filename)
            elif file_type == 'txt':
                return self._parse_txt(content, filename)
            else:
                raise ParseError(f"未知的文件类型: {file_type}")
        except Exception as e:
            logger.error(f"解析文件失败 {filename}: {e}")
            raise ParseError(f"解析失败: {str(e)}")
    
    def _parse_pdf(self, content: bytes, filename: str) -> ParsedDocument:
        """解析 PDF - 优先使用 pdfplumber，回退到 pypdf"""
        text = ""
        page_count = 0
        
        # 尝试 pdfplumber（中文支持更好）
        if self._pdfplumber_available:
            try:
                import pdfplumber
                import io
                
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    page_count = len(pdf.pages)
                    for i, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                # 强制 UTF-8 编码处理
                                page_text = page_text.encode('utf-8', errors='ignore').decode('utf-8')
                                text += f"\n--- Page {i+1} ---\n" + page_text
                        except Exception as e:
                            logger.warning(f"第 {i+1} 页解析失败: {e}")
                
                if text.strip():
                    logger.info(f"pdfplumber 解析成功: {filename}, {page_count} 页")
                    return ParsedDocument(
                        text=text.strip(),
                        filename=filename,
                        file_type='pdf',
                        page_count=page_count
                    )
            except Exception as e:
                logger.warning(f"pdfplumber 解析失败，尝试 pypdf: {e}")
        
        # 回退到 pypdf
        if self._pypdf_available:
            try:
                from pypdf import PdfReader
                import io
                
                reader = PdfReader(io.BytesIO(content))
                page_count = len(reader.pages)
                
                for i, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {i+1} ---\n" + page_text
                    except Exception as e:
                        logger.warning(f"第 {i+1} 页解析失败: {e}")
                
                if text.strip():
                    logger.info(f"pypdf 解析成功: {filename}, {page_count} 页")
                    return ParsedDocument(
                        text=text.strip(),
                        filename=filename,
                        file_type='pdf',
                        page_count=page_count
                    )
            except Exception as e:
                logger.error(f"pypdf 解析失败: {e}")
        
        # 都失败了
        if not self._pdfplumber_available and not self._pypdf_available:
            raise ParseError("PDF 解析库未安装，请运行: pip install pdfplumber pypdf")
        
        raise ParseError("PDF 解析失败，可能是扫描件或加密 PDF")
    
    def _parse_docx(self, content: bytes, filename: str) -> ParsedDocument:
        """解析 DOCX"""
        if not self._docx_available:
            raise ParseError("DOCX 解析库未安装，请运行: pip install python-docx")
        
        try:
            from docx import Document
            import io
            
            doc = Document(io.BytesIO(content))
            
            # 提取段落
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)
            
            text = '\n\n'.join(paragraphs)
            
            if not text.strip():
                raise ParseError("DOCX 文件内容为空")
            
            logger.info(f"DOCX 解析成功: {filename}, {len(paragraphs)} 段落")
            return ParsedDocument(
                text=text,
                filename=filename,
                file_type='docx',
                page_count=len(paragraphs) // 20 + 1  # 估算页数
            )
            
        except Exception as e:
            logger.error(f"DOCX 解析失败: {e}")
            raise ParseError(f"DOCX 解析失败: {str(e)}")
    
    def _parse_txt(self, content: bytes, filename: str) -> ParsedDocument:
        """解析 TXT - 自动检测编码"""
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        text = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                text = content.decode(encoding)
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            # 最后尝试，忽略错误
            text = content.decode('utf-8', errors='ignore')
            used_encoding = 'utf-8 (with ignore)'
        
        # 清理文本
        text = self._clean_text(text)
        
        if not text.strip():
            raise ParseError("TXT 文件内容为空")
        
        # 估算行数和页数
        lines = text.split('\n')
        page_count = len(lines) // 40 + 1
        
        logger.info(f"TXT 解析成功: {filename}, 编码={used_encoding}, {len(lines)} 行")
        return ParsedDocument(
            text=text,
            filename=filename,
            file_type='txt',
            page_count=page_count,
            metadata={'encoding': used_encoding}
        )
    
    def _clean_text(self, text: str) -> str:
        """清理文本"""
        # 移除多余的空白字符
        text = re.sub(r'\n{3,}', '\n\n', text)  # 多个换行合并
        text = re.sub(r'[ \t]+', ' ', text)  # 多个空格合并
        text = re.sub(r'\n[ \t]+', '\n', text)  # 行首空格移除
        return text.strip()


# 便捷函数
def parse_document(file: Union[bytes, BinaryIO], content_type: str, filename: str) -> ParsedDocument:
    """便捷解析函数"""
    parser = DocumentParser()
    return parser.parse(file, content_type, filename)
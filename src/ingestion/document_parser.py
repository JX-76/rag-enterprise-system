"""
Document Parser - 文档解析模块

支持格式:
- PDF (.pdf) - 使用pypdf
- Word (.docx) - 使用python-docx
- Markdown (.md) - 原生解析
- Text (.txt) - 原生解析

特性:
- 语义分块（按标题、段落）
- 重叠分块（保持上下文连续性）
- 元数据提取（文件名、页码、章节）
"""
import re
import os
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import logging

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """文档分块"""
    content: str
    source: str  # 源文件路径
    chunk_index: int  # 块序号
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def __post_init__(self):
        # 清理内容
        self.content = self._clean_text(self.content)
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """清理文本"""
        # 移除多余空白
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        return text.strip()
    
    def to_dict(self) -> Dict[str, Any]:
        """转为字典"""
        return {
            "content": self.content,
            "source": self.source,
            "chunk_index": self.chunk_index,
            "metadata": self.metadata,
            "char_count": len(self.content)
        }


@dataclass
class ParsedDocument:
    """解析后的文档"""
    source: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def __len__(self) -> int:
        return len(self.chunks)
    
    def get_total_chars(self) -> int:
        """获取总字符数"""
        return sum(len(c.content) for c in self.chunks)


class BaseParser(ABC):
    """解析器基类"""
    
    SUPPORTED_EXTENSIONS: List[str] = []
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 128,
        respect_semantic_boundaries: bool = True
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.respect_semantic_boundaries = respect_semantic_boundaries
    
    @abstractmethod
    def parse_file(self, file_path: str) -> ParsedDocument:
        """解析文件"""
        pass
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否支持该文件"""
        ext = Path(file_path).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS
    
    def chunk_text(
        self,
        text: str,
        source: str,
        base_metadata: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
        """
        将文本分块
        
        策略:
        1. 如果respect_semantic_boundaries=True，先按语义边界切分
        2. 对每个语义块再进行滑动窗口切分
        3. 保持重叠以保持上下文
        """
        chunks = []
        base_metadata = base_metadata or {}
        
        if self.respect_semantic_boundaries:
            # 按语义边界预切分
            semantic_units = self._split_by_semantic_boundaries(text)
        else:
            semantic_units = [text]
        
        chunk_index = 0
        for unit in semantic_units:
            # 如果语义单元比目标块大小小，直接作为一个块
            if len(unit) <= self.chunk_size:
                chunks.append(DocumentChunk(
                    content=unit,
                    source=source,
                    chunk_index=chunk_index,
                    metadata={**base_metadata, "semantic_unit": True}
                ))
                chunk_index += 1
            else:
                # 需要滑动窗口切分
                sub_chunks = self._sliding_window_split(unit)
                for sub_chunk in sub_chunks:
                    chunks.append(DocumentChunk(
                        content=sub_chunk,
                        source=source,
                        chunk_index=chunk_index,
                        metadata={**base_metadata, "semantic_unit": False}
                    ))
                    chunk_index += 1
        
        return chunks
    
    def _split_by_semantic_boundaries(self, text: str) -> List[str]:
        """
        按语义边界切分
        
        优先级:
        1. 标题 (## ###)
        2. 段落 (\n\n)
        3. 句子 (。！？)
        """
        units = []
        
        # 1. 按标题切分
        heading_pattern = r'(?:^|\n)(#{1,6}\s+.+?)(?=\n#{1,6}\s|\Z)'
        matches = list(re.finditer(heading_pattern, text, re.MULTILINE | re.DOTALL))
        
        if matches:
            # 有标题，按标题切分
            last_end = 0
            for match in matches:
                # 标题前的内容
                if match.start() > last_end:
                    pre_content = text[last_end:match.start()].strip()
                    if pre_content:
                        units.append(pre_content)
                
                # 标题及内容
                heading_end = text.find('\n#', match.end())
                if heading_end == -1:
                    heading_end = len(text)
                
                section = text[match.start():heading_end].strip()
                if section:
                    units.append(section)
                
                last_end = heading_end
            
            # 最后一部分
            if last_end < len(text):
                final = text[last_end:].strip()
                if final:
                    units.append(final)
        else:
            # 2. 按段落切分
            paragraphs = re.split(r'\n{2,}', text)
            for para in paragraphs:
                para = para.strip()
                if para:
                    units.append(para)
        
        return units if units else [text]
    
    def _sliding_window_split(self, text: str) -> List[str]:
        """滑动窗口切分"""
        chunks = []
        start = 0
        text_len = len(text)
        
        while start < text_len:
            # 计算结束位置
            end = min(start + self.chunk_size, text_len)
            
            # 如果不是最后一块，尝试在句子边界切分
            if end < text_len and self.respect_semantic_boundaries:
                # 向后找最近的句子结束符
                search_end = min(end + 50, text_len)
                sentence_end = self._find_sentence_end(text, end, search_end)
                if sentence_end > end:
                    end = sentence_end
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # 滑动窗口
            start = end - self.chunk_overlap
            if start <= 0:
                start = end
        
        return chunks
    
    def _find_sentence_end(self, text: str, start: int, end: int) -> int:
        """在范围内找句子结束位置"""
        sentence_chars = '.。！？\n'
        for i in range(start, end):
            if i < len(text) and text[i] in sentence_chars:
                return i + 1
        return start


class PDFParser(BaseParser):
    """PDF解析器"""
    
    SUPPORTED_EXTENSIONS = ['.pdf']
    
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查依赖"""
        try:
            import pypdf
            self._has_pypdf = True
        except ImportError:
            self._has_pypdf = False
            logger.warning("pypdf not installed. PDF parsing will be disabled.")
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否支持该文件"""
        if not self._has_pypdf:
            return False
        ext = Path(file_path).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS
    
    def parse_file(self, file_path: str) -> ParsedDocument:
        """解析PDF文件"""
        if not self._has_pypdf:
            raise ImportError("pypdf not installed. Install: pip install pypdf")
        
        from pypdf import PdfReader
        
        logger.info(f"Parsing PDF: {file_path}")
        
        reader = PdfReader(file_path)
        full_text = ""
        page_metadata = []
        
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                full_text += f"\n\n--- Page {i+1} ---\n\n{text}"
                page_metadata.append({
                    "page": i + 1,
                    "text_length": len(text)
                })
        
        # 分块
        chunks = self.chunk_text(
            full_text,
            source=file_path,
            base_metadata={
                "file_type": "pdf",
                "total_pages": len(reader.pages),
                "pages": page_metadata
            }
        )
        
        return ParsedDocument(
            source=file_path,
            chunks=chunks,
            metadata={
                "file_type": "pdf",
                "total_pages": len(reader.pages),
                "parser": "pypdf"
            }
        )


class MarkdownParser(BaseParser):
    """Markdown解析器"""
    
    SUPPORTED_EXTENSIONS = ['.md', '.markdown']
    
    def parse_file(self, file_path: str) -> ParsedDocument:
        """解析Markdown文件"""
        logger.info(f"Parsing Markdown: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取元数据（YAML frontmatter）
        metadata = self._extract_metadata(content)
        
        # 移除frontmatter后的内容
        text = self._remove_frontmatter(content)
        
        chunks = self.chunk_text(
            text,
            source=file_path,
            base_metadata={
                "file_type": "markdown",
                **metadata
            }
        )
        
        return ParsedDocument(
            source=file_path,
            chunks=chunks,
            metadata={
                "file_type": "markdown",
                "parser": "native",
                "frontmatter": metadata
            }
        )
    
    def _extract_metadata(self, content: str) -> Dict[str, Any]:
        """提取YAML frontmatter"""
        if content.startswith('---'):
            end = content.find('---', 3)
            if end != -1:
                yaml_content = content[3:end].strip()
                # 简单解析key: value
                metadata = {}
                for line in yaml_content.split('\n'):
                    if ':' in line:
                        key, value = line.split(':', 1)
                        metadata[key.strip()] = value.strip()
                return metadata
        return {}
    
    def _remove_frontmatter(self, content: str) -> str:
        """移除frontmatter"""
        if content.startswith('---'):
            end = content.find('---', 3)
            if end != -1:
                return content[end + 3:].strip()
        return content


class TextParser(BaseParser):
    """纯文本解析器"""
    
    SUPPORTED_EXTENSIONS = ['.txt', '.text', '.rst']
    
    def parse_file(self, file_path: str) -> ParsedDocument:
        """解析文本文件"""
        logger.info(f"Parsing Text: {file_path}")
        
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        content = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError(f"Cannot decode file: {file_path}")
        
        chunks = self.chunk_text(
            content,
            source=file_path,
            base_metadata={"file_type": "text", "encoding": encoding}
        )
        
        return ParsedDocument(
            source=file_path,
            chunks=chunks,
            metadata={
                "file_type": "text",
                "parser": "native",
                "encoding": encoding
            }
        )


class DocxParser(BaseParser):
    """Word文档解析器"""
    
    SUPPORTED_EXTENSIONS = ['.docx']
    
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查依赖"""
        try:
            import docx
            self._has_docx = True
        except ImportError:
            self._has_docx = False
            logger.warning("python-docx not installed. DOCX parsing will be disabled.")
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否支持该文件"""
        if not self._has_docx:
            return False
        ext = Path(file_path).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS
    
    def parse_file(self, file_path: str) -> ParsedDocument:
        """解析DOCX文件"""
        if not self._has_docx:
            raise ImportError("python-docx not installed. Install: pip install python-docx")
        
        from docx import Document
        
        logger.info(f"Parsing DOCX: {file_path}")
        
        doc = Document(file_path)
        
        # 提取文本
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        text = '\n\n'.join(full_text)
        
        chunks = self.chunk_text(
            text,
            source=file_path,
            base_metadata={
                "file_type": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
        )
        
        return ParsedDocument(
            source=file_path,
            chunks=chunks,
            metadata={
                "file_type": "docx",
                "parser": "python-docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
        )


class DocumentParser:
    """文档解析器统一入口"""
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 128,
        respect_semantic_boundaries: bool = True
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.respect_semantic_boundaries = respect_semantic_boundaries
        
        # 初始化所有解析器
        self.parsers: List[BaseParser] = [
            PDFParser(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                respect_semantic_boundaries=respect_semantic_boundaries
            ),
            MarkdownParser(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                respect_semantic_boundaries=respect_semantic_boundaries
            ),
            TextParser(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                respect_semantic_boundaries=respect_semantic_boundaries
            ),
            DocxParser(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                respect_semantic_boundaries=respect_semantic_boundaries
            )
        ]
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析单个文件"""
        file_path = str(file_path)
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser.parse_file(file_path)
        
        raise ValueError(
            f"Unsupported file format: {Path(file_path).suffix}. "
            f"Supported: {self.get_supported_extensions()}"
        )
    
    def parse_directory(
        self,
        directory: str,
        recursive: bool = True
    ) -> List[ParsedDocument]:
        """解析目录下所有文件"""
        documents = []
        path = Path(directory)
        
        pattern = '**/*' if recursive else '*'
        
        for file_path in path.glob(pattern):
            if file_path.is_file():
                try:
                    doc = self.parse(str(file_path))
                    documents.append(doc)
                    logger.info(f"Parsed: {file_path} -> {len(doc)} chunks")
                except Exception as e:
                    logger.warning(f"Failed to parse {file_path}: {e}")
        
        return documents
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        extensions = []
        for parser in self.parsers:
            extensions.extend(parser.SUPPORTED_EXTENSIONS)
        return extensions


# 便捷函数
def parse_document(
    file_path: str,
    chunk_size: int = 512,
    chunk_overlap: int = 128
) -> ParsedDocument:
    """便捷函数：解析单个文档"""
    parser = DocumentParser(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    return parser.parse(file_path)


def parse_directory(
    directory: str,
    chunk_size: int = 512,
    chunk_overlap: int = 128,
    recursive: bool = True
) -> List[ParsedDocument]:
    """便捷函数：解析目录"""
    parser = DocumentParser(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    return parser.parse_directory(directory, recursive)